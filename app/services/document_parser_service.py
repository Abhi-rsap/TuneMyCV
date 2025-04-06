import io
from docx import Document as DocxDocument
from app.models.resume_model import Resume, ContactInfo
from app.models.document_model import Document
from pymongo import MongoClient
from datetime import datetime, timezone
from bson import ObjectId
import re

# Set up MongoDB client
client = MongoClient("mongodb://localhost:27017/")
db = client.resume_bot
collection = db.documents

def parse_word_document(file_data: bytes) -> Resume:
    doc = DocxDocument(io.BytesIO(file_data))
    
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    contact_info = ContactInfo()
    contact_info.name = paragraphs[0] if len(paragraphs) > 0 else ""
    info = paragraphs[1].split("•")
    for i in range(len(info)):
        info[i] = info[i].strip()
        if "@" in info[i]:
            contact_info.email = info[i]
        elif "linkedin" in info[i].lower():
            contact_info.linkedin = info[i]
        elif "github" in info[i].lower():
            contact_info.github = info[i]
        elif re.match(r"^\+?\d{10,15}$", info[i]):
            contact_info.phone = info[i]


    resume = Resume(
        contact_information=contact_info,
        summary="Hi this is Abhi",
        education=[],
        work_experience=[],
        skills=[],
        additonal_sections={}
    )
    return resume

def save_document_to_db(title: str, resume_data: Resume, filename: str) -> dict:
    document = Document(
        title=title or filename,
        content=resume_data,
        uploaded_at=datetime.now(timezone.utc)
    )
    
    doc_dict = document.dict(by_alias=True)
    
    result = collection.insert_one(doc_dict)
    doc_dict["_id"] = str(result.inserted_id)
    return doc_dict

def handle_document_upload(file_data: bytes, filename: str) -> dict:
    resume_data = parse_word_document(file_data)
    return save_document_to_db(resume_data.contact_information.full_name or "Untitled", resume_data, filename)
